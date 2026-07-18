"""文档解析：pdf / docx / txt / md -> 纯文本"""
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}


def load_text(file_path: str, file_type: str) -> str:
    path = Path(file_path)
    if file_type == "pdf":
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if file_type == "docx":
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格内容按行拼接，避免制度类文档中表格信息丢失
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    # txt / md：优先 utf-8，兼容 Windows 记事本的 gbk
    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("gbk", errors="ignore")
